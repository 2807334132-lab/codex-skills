import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def fmt_time(value):
    seconds = int(round(float(value)))
    return f"{seconds // 60:02d}:{seconds % 60:02d}"


def set_font(run, size, color=None, bold=None):
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def main():
    ap = argparse.ArgumentParser(description="Build a timestamped transcript DOCX from corrected JSON")
    ap.add_argument("input_json", type=Path)
    ap.add_argument("output_docx", type=Path)
    ap.add_argument("--title", default="视频音频逐字稿")
    ap.add_argument("--source", default="")
    ap.add_argument("--duration", default="")
    args = ap.parse_args()
    rows = json.loads(args.input_json.read_text(encoding="utf-8"))
    if not isinstance(rows, list) or not rows:
        raise ValueError("input JSON must be a non-empty list")
    for i, row in enumerate(rows):
        if not all(k in row for k in ("start", "end", "text")) or not str(row["text"]).strip():
            raise ValueError(f"invalid transcript item at index {i}")

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Microsoft YaHei", Pt(11)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after, normal.paragraph_format.line_spacing = Pt(7), 1.15

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(args.title), 24, "0B2545", True)
    if args.source:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        set_font(p.add_run(f"来源：{args.source}"), 10.5, "4B5563")
    if args.duration:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        set_font(p.add_run(f"时长：{args.duration}"), 10.5, "4B5563")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    set_font(p.add_run("说明：按音频逐句整理；时间码为近似值，便于回听核对。"), 9.5, "5A626C")

    for row in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(0), Pt(7)
        p.paragraph_format.line_spacing = 1.15
        set_font(p.add_run(f"[{fmt_time(row['start'])}]  "), 9.5, "2E74B5", True)
        set_font(p.add_run(str(row["text"]).strip()), 11, "1F2937")

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("第 "), 8.5, "6B7280")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)
    set_font(footer.add_run(" 页"), 8.5, "6B7280")

    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = args.title
    doc.core_properties.author = ""
    doc.save(args.output_docx)
    print(args.output_docx.resolve())


if __name__ == "__main__":
    main()

